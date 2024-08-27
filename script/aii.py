import cv2
import pytesseract
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
img = cv2.imread("your_image.jpg")
text = pytesseract.image_to_string(img)
from docx import Document

def create_word_document(text):
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run(text)
    document.save("output.docx")

create_word_document(text)
