import cv2
import pytesseract
from docx import Document
# Set the path to the Tesseract executable (change this to your installation path)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Read the image
img = cv2.imread(".ocr/dillated.jpg")

# Extract text from the image
text = pytesseract.image_to_string(img)
from docx import Document

def create_word_document(text):
    document = Document()
    paragraph = document.add_paragraph(text)
    document.save("output.docx")

create_word_document(text)


from docx.shared import Inches

def add_image_to_word(image_path):
    document = Document("output.docx")
    paragraph = document.add_paragraph()
    paragraph.add_run("Image from OpenCV:")
    paragraph.add_picture(image_path, width=Inches(6.5))
    document.save("output.docx")

add_image_to_word("your_image.jpg")
